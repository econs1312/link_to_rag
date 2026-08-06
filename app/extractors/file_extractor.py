"""
File Extractor: handles PDF, DOCX, and image files (with OCR).

Security & Validation (skill 10):
  - Uses python-magic to verify real MIME type via magic bytes (not extension).
  - Rejects files whose real content doesn't match supported types.

Extraction strategy:
  - PDF:    pymupdf (fitz) for native text extraction.
            Falls back to Tesseract OCR page-by-page when a page has < 100 chars
            of native text (scanned PDF detection per skill 10 §3).
  - DOCX:   python-docx for structured paragraph + table extraction.
  - Images: OpenAI Vision API (if OPENAI_API_KEY configured) for best accuracy,
            then falls back to Tesseract OCR for local processing.

Memory safety (skill 10 §4):
  - All temp files cleaned in try...finally blocks.
  - Pixmaps explicitly freed after OCR render.
  - PIL Images use context managers.
"""

import io
import os
import asyncio
import base64
import tempfile
from typing import Optional

from app.schemas.ingestion import ExtractedContent
from app.core.logging import logger
from app.core.config import settings


# ─────────────────────────────────────────────────────────────────────────────
# MIME / extension helpers
# ─────────────────────────────────────────────────────────────────────────────

SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc_legacy",
    "image/png": "image",
    "image/jpeg": "image",
    "image/jpg": "image",
    "image/tiff": "image",
    "image/bmp": "image",
    "image/gif": "image",
    "image/webp": "image",
}

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc_legacy",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".tiff": "image",
    ".tif": "image",
    ".bmp": "image",
    ".gif": "image",
    ".webp": "image",
}

# Minimum characters per PDF page to consider native text valid (skill 10 §3)
_OCR_FALLBACK_THRESHOLD = 100


def detect_file_type(filename: str, content_type: Optional[str] = None) -> Optional[str]:
    """Detect file type from MIME type or file extension."""
    if content_type and content_type in SUPPORTED_MIME_TYPES:
        return SUPPORTED_MIME_TYPES[content_type]
    ext = os.path.splitext(filename.lower())[1]
    return SUPPORTED_EXTENSIONS.get(ext)


def validate_real_mime_type(file_bytes: bytes, filename: str) -> str:
    """
    Verify real file type using magic bytes (skill 10 §1).

    Returns the resolved file type string ('pdf', 'docx', 'image').
    Raises ValueError if the file type is not supported or is suspicious.
    """
    try:
        import magic
        real_mime = magic.from_buffer(file_bytes, mime=True)
    except ImportError:
        logger.warning(
            "python-magic not available, falling back to extension-based detection",
            filename=filename,
        )
        file_type = detect_file_type(filename)
        if not file_type:
            raise ValueError(f"Unsupported file type: '{filename}'")
        return file_type
    except Exception as exc:
        logger.warning(
            "Magic bytes detection failed, falling back to extension",
            filename=filename,
            error=str(exc),
        )
        file_type = detect_file_type(filename)
        if not file_type:
            raise ValueError(f"Unsupported file type: '{filename}'")
        return file_type

    # Map real MIME to our file type
    resolved_type = SUPPORTED_MIME_TYPES.get(real_mime)
    if not resolved_type:
        raise ValueError(
            f"Ficheiro rejeitado: o tipo real do ficheiro é '{real_mime}' "
            f"(detectado por magic bytes), que não é suportado. "
            f"Ficheiro: '{filename}'. "
            f"Tipos suportados: PDF, DOCX, PNG, JPG, TIFF, BMP, GIF, WEBP."
        )

    # Log warning if declared type differs from real type
    declared_type = detect_file_type(filename)
    if declared_type and declared_type != resolved_type:
        logger.warning(
            "File real MIME type differs from declared extension",
            filename=filename,
            declared_type=declared_type,
            real_mime=real_mime,
            resolved_type=resolved_type,
        )

    return resolved_type


# ─────────────────────────────────────────────────────────────────────────────
# Main extraction dispatcher
# ─────────────────────────────────────────────────────────────────────────────

async def extract_file_content(
    file_bytes: bytes,
    filename: str,
    content_type: Optional[str] = None,
) -> ExtractedContent:
    """
    Entry point: validates real MIME type, then dispatches extraction.
    Returns ExtractedContent compatible with the existing ingestion pipeline.
    """
    # 1. Validate real MIME type using magic bytes (skill 10 §1)
    file_type = validate_real_mime_type(file_bytes, filename)

    logger.info("File extraction dispatched", filename=filename, file_type=file_type)

    loop = asyncio.get_event_loop()

    if file_type == "pdf":
        text, metadata = await loop.run_in_executor(None, _extract_pdf, file_bytes, filename)
    elif file_type in ("docx", "doc_legacy"):
        text, metadata = await loop.run_in_executor(None, _extract_docx, file_bytes, filename)
    elif file_type == "image":
        text, metadata = await _extract_image(file_bytes, filename)
    else:
        raise ValueError(f"Unknown file_type: {file_type}")

    return ExtractedContent(
        raw_text=text or f"[Nenhum texto extraído do arquivo: {filename}]",
        title=metadata.get("title", filename),
        author=metadata.get("author", "Arquivo Local"),
        metadata={
            "source": "file_upload",
            "file_type": file_type,
            "filename": filename,
            **metadata,
        },
        source_url=f"file://{filename}",
    )


# ─────────────────────────────────────────────────────────────────────────────
# PDF Extraction (pymupdf + Tesseract fallback)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """
    Extract text from PDF using pymupdf (skill 10 §2).
    Falls back to Tesseract OCR for pages with < 100 chars of native text (skill 10 §3).
    """
    import fitz  # pymupdf

    pages_text: list[str] = []
    ocr_pages = 0
    native_pages = 0

    doc = None
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        title = doc.metadata.get("title") or filename
        author = doc.metadata.get("author") or "PDF Document"
        num_pages = len(doc)

        for page_num, page in enumerate(doc, start=1):
            native_text = page.get_text().strip()

            if native_text and len(native_text) > _OCR_FALLBACK_THRESHOLD:
                # Page has a proper text layer (skill 10 §3: > 100 chars)
                pages_text.append(f"\n## Página {page_num}\n\n{native_text}")
                native_pages += 1
            else:
                # Image-only or sparse page — render to image and OCR (skill 10 §3)
                logger.debug(
                    "PDF page has insufficient text, applying Tesseract OCR",
                    page=page_num,
                    filename=filename,
                    native_chars=len(native_text) if native_text else 0,
                )
                pix = None
                try:
                    pix = page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("png")
                    ocr_text = _tesseract_ocr_from_bytes(img_bytes)
                    if ocr_text.strip():
                        pages_text.append(f"\n## Página {page_num} (OCR)\n\n{ocr_text}")
                        ocr_pages += 1
                finally:
                    # Explicitly free pixmap memory (skill 10 §4)
                    if pix is not None:
                        pix = None  # noqa: F841 — release reference for GC

    finally:
        # Ensure document handle is always closed (skill 10 §4)
        if doc is not None:
            doc.close()

    full_text = "\n".join(pages_text)
    return full_text, {
        "title": title,
        "author": author,
        "num_pages": num_pages,
        "native_pages": native_pages,
        "ocr_pages": ocr_pages,
        "extraction_method": "pymupdf" + ("+tesseract_ocr" if ocr_pages > 0 else ""),
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Extraction (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """Extract text from DOCX using python-docx, including paragraphs and tables."""
    from docx import Document as DocxDocument

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        doc = DocxDocument(tmp_path)
        parts: list[str] = []

        # Extract paragraphs (preserving headings)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name.lower() if para.style else ""
            if "heading 1" in style_name:
                parts.append(f"\n# {text}")
            elif "heading 2" in style_name:
                parts.append(f"\n## {text}")
            elif "heading 3" in style_name:
                parts.append(f"\n### {text}")
            else:
                parts.append(text)

        # Extract tables as Markdown
        for table_idx, table in enumerate(doc.tables, start=1):
            parts.append(f"\n### Tabela {table_idx}\n")
            rows_md: list[str] = []
            for row_i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows_md.append("| " + " | ".join(cells) + " |")
                if row_i == 0:
                    rows_md.append("|" + "|".join(["---"] * len(cells)) + "|")
            parts.append("\n".join(rows_md))

        # Core properties (author, title)
        props = doc.core_properties
        title = getattr(props, "title", None) or filename
        author = getattr(props, "author", None) or "DOCX Document"

        full_text = "\n\n".join(parts)
        return full_text, {
            "title": title,
            "author": author,
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables),
            "extraction_method": "python-docx",
        }
    finally:
        # Always clean up temp file (skill 10 §4)
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ─────────────────────────────────────────────────────────────────────────────
# Image OCR (OpenAI Vision API preferred, Tesseract fallback)
# ─────────────────────────────────────────────────────────────────────────────

async def _extract_image(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """OCR an image file. Uses OpenAI Vision API if key is available, else Tesseract."""
    if settings.OPENAI_API_KEY:
        logger.info("Using OpenAI Vision API for image OCR", filename=filename)
        try:
            text = await _openai_vision_ocr(file_bytes, filename)
            return text, {
                "title": filename,
                "author": "Image OCR",
                "extraction_method": "openai_vision",
            }
        except Exception as exc:
            logger.warning("OpenAI Vision API failed, falling back to Tesseract", error=str(exc), filename=filename)

    # Tesseract fallback
    logger.info("Using Tesseract OCR for image", filename=filename)
    loop = asyncio.get_event_loop()
    text = await loop.run_in_executor(None, _tesseract_ocr_from_bytes, file_bytes)
    return text, {
        "title": filename,
        "author": "Image OCR",
        "extraction_method": "tesseract",
    }


async def _openai_vision_ocr(file_bytes: bytes, filename: str) -> str:
    """Call OpenAI GPT-4o Vision to extract text from an image (superior to Tesseract for complex layouts)."""
    import httpx

    b64_image = base64.b64encode(file_bytes).decode("utf-8")
    ext = os.path.splitext(filename.lower())[1].lstrip(".")
    mime_map = {"jpg": "jpeg", "tif": "tiff"}
    mime_type = f"image/{mime_map.get(ext, ext)}"

    payload = {
        "model": "gpt-4o-mini",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extraia e transcreva todo o texto visível nesta imagem de forma fiel e completa. "
                            "Mantenha a estrutura e formatação do texto original (títulos, parágrafos, listas, tabelas). "
                            "Responda APENAS com o texto extraído, sem explicações adicionais."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{b64_image}", "detail": "high"},
                    },
                ],
            }
        ],
        "max_tokens": 4096,
    }

    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(
            "https://api.openai.com/v1/chat/completions",
            json=payload,
            headers={"Authorization": f"Bearer {settings.OPENAI_API_KEY}"},
        )
        resp.raise_for_status()
        data = resp.json()
        return data["choices"][0]["message"]["content"].strip()


def _tesseract_ocr_from_bytes(image_bytes: bytes) -> str:
    """Run Tesseract OCR on raw image bytes. Supports Portuguese + English."""
    try:
        import pytesseract
        from PIL import Image

        # Use context manager for PIL Image (skill 10 §4)
        with Image.open(io.BytesIO(image_bytes)) as image:
            # Use Portuguese + English for best coverage
            text = pytesseract.image_to_string(image, lang="por+eng", config="--oem 3 --psm 6")
            return text.strip()
    except Exception as exc:
        logger.error("Tesseract OCR failed", error=str(exc))
        return ""
